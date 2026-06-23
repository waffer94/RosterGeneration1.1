#!/usr/bin/env python3
"""
build_rosters.py
================
Generate F.A.S.T. Rescue course rosters (Word .docx) from an Excel export.

Usage:
    python build_rosters.py EXPORT.xlsx [--templates TEMPLATE_DIR] [--out OUTPUT_DIR]

For every data row in the export it:
  1. Detects the course type (In-Class / Blended / Recertification) from the Name column.
  2. Picks the matching blank template *by its layout*, not its filename
     (the In-Class / Blended template files are cross-named, so we identify by content).
  3. Fills the info table and the participant table, preserving all template formatting.
  4. Saves "<Instructor> - <Start date> - <Company>.docx".
"""

import argparse
import copy
import glob
import os
import re
import sys
from datetime import datetime, time as dtime

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------------------------------------------------------- #
# Field sanitization
# ----------------------------------------------------------------------------- #

def clean_company(customer: str) -> str:
    """'519817 CUPE Ontario' -> 'CUPE Ontario' (strip leading account id)."""
    return re.sub(r"^\s*\d+\s+", "", (customer or "").strip()).strip()


def extract_course_name(name_col: str) -> str:
    """
    Pull the clean course name out of the Name column, e.g.
    '519817 CUPE Ontario : 20260506 In-Class Standard First Aid & CPR/AED Level C
     Recertificatio Markham'  ->  'In-Class Standard First Aid & CPR/AED Level C Recertification'
    """
    raw = (name_col or "").strip()
    # drop the "<id customer> : " prefix
    after = raw.split(" : ", 1)[1] if " : " in raw else raw
    # drop a leading YYYYMMDD date token
    after = re.sub(r"^\s*\d{8}\s+", "", after).strip()

    # Anchor on "... Level <X>" and decide on the Recertification suffix from the tail.
    m = re.match(r"^(.*?Level\s+[A-Za-z])\b(.*)$", after)
    if m:
        base, tail = m.group(1).strip(), m.group(2)
        if "recert" in tail.lower():
            base += " Recertification"
        return base

    # Fallback (course name has no "Level X"): the word "Training" marks the
    # boundary before the trailing city, so cut there. Otherwise keep as-is.
    after = re.split(r"\bTraining\b", after, maxsplit=1)[0].strip()
    after = re.sub(r"\bRecertificatio(n)?\b", "Recertification", after, flags=re.I)
    return re.sub(r"\s+", " ", after).strip()


def extract_location(location: str, company: str, contact: str = "") -> str:
    """Strip a leading contact name and/or company name off the address."""
    loc = (location or "").strip()
    prefixes = [p for p in (contact, company) if p]
    again = True
    while again:
        again = False
        for pre in prefixes:
            if loc.lower().startswith(pre.lower()):
                loc = loc[len(pre):].lstrip(" ,-")
                again = True
    return re.sub(r"\s+", " ", loc).strip()


def extract_contact_name(contact_col: str) -> str:
    """'519817 CUPE Ontario : Olivia Kirby' -> 'Olivia Kirby'."""
    raw = (contact_col or "").strip()
    return raw.rsplit(" : ", 1)[-1].strip() if " : " in raw else raw


def format_contact(name: str, phone) -> str:
    """'Name – phone' (en-dash). Phone passed through verbatim."""
    phone = "" if phone is None else str(phone).strip()
    name = (name or "").strip()
    if name and phone:
        return f"{name} \u2013 {phone}"
    return name or phone


def _as_dt(v):
    if isinstance(v, datetime):
        return v
    if isinstance(v, str) and v.strip():
        for fmt in ("%m/%d/%Y", "%Y-%m-%d", "%Y%m%d"):
            try:
                return datetime.strptime(v.strip(), fmt)
            except ValueError:
                pass
    return None


def _fmt_single(d) -> str:
    return f"{d:%B} {d.day}, {d.year}" if d else ""


def _first_date(start, other, end):
    """Earliest valid session date (drops dates before the start, e.g. typos)."""
    s = _as_dt(start)
    collected = [d for d in (_as_dt(start), _as_dt(other), _as_dt(end)) if d]
    if not collected:
        return None
    anchor = s or min(collected)
    valid = [d for d in collected if d.date() >= anchor.date()]
    return min(valid) if valid else anchor


def format_dates(start, other, end) -> str:
    """
    List the actual session dates from Course Date / Other Date / Course End Date.
      single day       -> 'May 2, 2026'
      two days         -> 'May 28 & 29, 2026'
      several sessions -> 'May 2, 3 & 5, 2026'   (month repeated only when it changes)
    Dates earlier than the start (e.g. a typo'd year) are dropped.
    """
    s = _as_dt(start)
    collected = [d for d in (_as_dt(start), _as_dt(other), _as_dt(end)) if d]
    if not collected:
        return ""
    anchor = s or min(collected)
    uniq = sorted({d.date() for d in collected if d.date() >= anchor.date()})
    if len(uniq) == 1:
        d = uniq[0]
        return f"{d.strftime('%B')} {d.day}, {d.year}"

    multi_year = len({d.year for d in uniq}) > 1
    pieces, prev_month = [], None
    for d in uniq:
        if multi_year:
            pieces.append(f"{d.strftime('%B')} {d.day}, {d.year}")
        elif d.month != prev_month:
            pieces.append(f"{d.strftime('%B')} {d.day}")
        else:
            pieces.append(str(d.day))
        prev_month = d.month
    body = pieces[0] if len(pieces) == 1 else ", ".join(pieces[:-1]) + " & " + pieces[-1]
    return body if multi_year else f"{body}, {uniq[0].year}"


def _as_time(v):
    if isinstance(v, dtime):
        return v
    if isinstance(v, datetime):
        return v.time()
    if isinstance(v, str) and v.strip():
        for fmt in ("%H:%M", "%I:%M %p", "%I:%M%p"):
            try:
                return datetime.strptime(v.strip(), fmt).time()
            except ValueError:
                pass
    return None


def _fmt_t(t) -> str:
    # 8:30 am  (no leading zero, lowercase meridiem)
    return f"{(t.hour % 12) or 12}:{t.minute:02d} {'am' if t.hour < 12 else 'pm'}"


def format_time(t_from, t_to) -> str:
    a, b = _as_time(t_from), _as_time(t_to)
    if a and b:
        return f"{_fmt_t(a)} to {_fmt_t(b)}"
    return _fmt_t(a) if a else ""


BULLET_RE = re.compile(r"^\s*[\u2022\u00b7\u25aa\u2023\u25e6\*\-]\s+")
EMAILISH = re.compile(r"[^\s@]+@[^\s@]+")


def _is_email_tok(tok: str) -> bool:
    return bool(EMAILISH.search(tok))


def _is_header(line: str) -> bool:
    """Group/section header lines that are not participants."""
    l = line.strip()
    if not l:
        return False
    if re.search(r"group\s*\(", l, re.I):
        return True
    if l.endswith(":"):
        return True
    return False


def _invert_comma(name: str) -> str:
    """'Hwang, Mikaela' -> 'Mikaela Hwang'."""
    if name.count(",") == 1:
        last, first = [p.strip() for p in name.split(",")]
        if last and first:
            return f"{first} {last}"
    return name


def _finalize(raw_name: str):
    """Collapse spaces, split a trailing '- note', invert 'Last, First'."""
    parts = re.split(r"-\s+", raw_name, maxsplit=1)
    name = re.sub(r"\s+", " ", parts[0]).strip(" \t-\u2022\u00b7")
    name = re.sub(r"^\d+\s*[\.\):]?\s+", "", name)   # drop any leading "1." / "1)" / "1 "
    note = re.sub(r"\s+", " ", parts[1]).strip() if len(parts) > 1 else ""
    name = _invert_comma(name)
    return name, note


def _names_from_block(text: str):
    """
    Auto-detect the participant format of one text block and return raw name strings.
      * has emails   -> accumulate lines until an email terminates a person
      * else bullets -> each bullet marks a new person
      * else         -> one person per line
    Leading bullets, emails, blank lines and group headers are handled in every mode.
    """
    lines = [l.strip() for l in re.split(r"[\r\n]+", text.replace("_x000D_", ""))]
    has_email = any("@" in l for l in lines)
    has_bullet = any(BULLET_RE.match(l) for l in lines)
    people, cur = [], []

    def flush():
        if cur:
            people.append(" ".join(cur))
            cur.clear()

    if has_email:
        for l in lines:
            if not l:
                flush(); continue
            if _is_header(l):
                flush(); continue
            l = BULLET_RE.sub("", l)
            toks = l.split()
            cur.extend(t for t in toks if not _is_email_tok(t))
            if any(_is_email_tok(t) for t in toks):
                flush()
        flush()
    elif has_bullet:
        for l in lines:
            if not l:
                flush(); continue
            if _is_header(l):
                flush(); continue
            if BULLET_RE.match(l):
                flush()
                l = BULLET_RE.sub("", l)
            cur.extend(t for t in l.split() if not _is_email_tok(t))
        flush()
    else:
        for l in lines:
            if not l or _is_header(l):
                continue
            people.append(BULLET_RE.sub("", l))
    return people


def parse_participants(part_col, recert_col, sort_alpha: bool = True):
    """Returns a list of (name, note) tuples, format auto-detected per block."""
    rows = []
    for blob in (part_col, recert_col):
        if not blob:
            continue
        for raw in _names_from_block(str(blob)):
            name, note = _finalize(raw)
            if name:
                rows.append((name, note))
    if sort_alpha:
        rows.sort(key=lambda r: r[0].lower())
    return rows


# ----------------------------------------------------------------------------- #
# Template handling
# ----------------------------------------------------------------------------- #

def detect_course_type(name_col: str) -> str:
    low = (name_col or "").lower()
    if "working at height" in low:          # 'Working at Heights' (+ variants)
        return "working_at_heights"
    if "recert" in low:
        return "recert"
    if "blended" in low:
        return "blended"
    if "in-class" in low or "in class" in low:
        return "in-class"
    return "general"  # anything unrecognised -> General template


def classify_template(path: str) -> str:
    """Identify a template. New templates are matched by filename; the First Aid
    trio is matched by content (their In-Class/Blended filenames are unreliable)."""
    base = re.sub(r"[_\-]+", " ", os.path.basename(path).lower())
    if "working at height" in base:
        return "working_at_heights"
    if "general" in base:
        return "general"
    d = Document(path)
    header = [c.text.strip().lower() for c in d.tables[1].rows[0].cells]
    if any("online part" in h for h in header):
        return "blended"
    material = ""
    for row in d.tables[0].rows:
        if row.cells[0].text.strip().lower() == "course material":
            material = row.cells[1].text.lower()
            break
    return "in-class" if "manual" in material else "recert"


def load_template_map(template_dir: str) -> dict:
    mapping = {}
    for path in glob.glob(os.path.join(template_dir, "*.docx")):
        if "template" not in os.path.basename(path).lower():
            continue
        try:
            mapping[classify_template(path)] = path
        except Exception as exc:                       # noqa: BLE001
            print(f"  ! could not classify {os.path.basename(path)}: {exc}",
                  file=sys.stderr)
    return mapping


# ----------------------------------------------------------------------------- #
# Word writing helpers
# ----------------------------------------------------------------------------- #

def _strip_list_p(p_el):
    """Remove auto-list numbering / list indentation from a <w:p> element."""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        return
    for tag in ("w:numPr", "w:ind"):
        el = pPr.find(qn(tag))
        if el is not None:
            pPr.remove(el)
    ps = pPr.find(qn("w:pStyle"))
    if ps is not None and ps.get(qn("w:val"), "").lower().startswith("list"):
        pPr.remove(ps)


def set_cell(cell, text, bold=False, font="Times New Roman", size=12, center=False):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    _strip_list_p(p._p)
    if center:
        p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    return run


def strip_table_numbering(table):
    """Remove auto-list numbering from every data cell (keeps header row)."""
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                _strip_list_p(p._p)


def _blank_row_like(table, ref_tr):
    new_tr = copy.deepcopy(ref_tr)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for c in row.cells:
        for p in c.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
    return row


def _set_tc_width(tc, w):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(int(w)))


def add_number_column(table, header="Sr. No.", width=850):
    """Prepend a 'Sr. No.' column and number every data row 1..N."""
    grid = table._tbl.find(qn("w:tblGrid"))
    gcols = grid.findall(qn("w:gridCol"))
    name_w = int(gcols[0].get(qn("w:w")))
    new_name_w = max(name_w - width, 600)
    gcols[0].set(qn("w:w"), str(new_name_w))
    new_gc = OxmlElement("w:gridCol")
    new_gc.set(qn("w:w"), str(width))
    grid.insert(0, new_gc)

    for row in table.rows:                       # clone each row's own first cell
        first_tc = row._tr.findall(qn("w:tc"))[0]
        new_tc = copy.deepcopy(first_tc)
        paras = new_tc.findall(qn("w:p"))
        for extra in paras[1:]:                  # keep a single, clean paragraph
            new_tc.remove(extra)
        p0 = new_tc.find(qn("w:p"))
        for r in p0.findall(qn("w:r")):
            p0.remove(r)
        _strip_list_p(p0)
        _set_tc_width(new_tc, width)
        _set_tc_width(first_tc, new_name_w)
        first_tc.addprevious(new_tc)

    n = 0
    for ri, row in enumerate(table.rows):        # write header then 1..N
        if ri == 0:
            set_cell(row.cells[0], header, bold=True, center=True)
        else:
            n += 1
            set_cell(row.cells[0], str(n), center=True)


def fill_info_table(doc, info: dict):
    table = doc.tables[0]
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label in info:
            set_cell(row.cells[1], info[label])


def ensure_data_rows(table, target):
    """Make the participant table have exactly `target` data rows (excl. header)."""
    data = table.rows[1:]
    cur = len(data)
    if cur < target:
        ref_tr = data[-1]._tr if data else table.rows[-1]._tr
        for _ in range(target - cur):
            _blank_row_like(table, ref_tr)
    elif cur > target:
        for _ in range(cur - target):
            tr = table.rows[-1]._tr
            tr.getparent().remove(tr)


def fill_participants(doc, participants):
    table = doc.tables[1]
    header = [c.text.strip().lower() for c in table.rows[0].cells]
    note_idx = next((i for i, h in enumerate(header) if "note" in h), len(header) - 1)
    target = max(20, len(participants) + 5)   # >=20 rows, always >=5 blank rows
    ensure_data_rows(table, target)
    data_rows = table.rows[1:]
    for i, (name, note) in enumerate(participants):
        set_cell(data_rows[i].cells[0], name)
        if note:
            set_cell(data_rows[i].cells[note_idx], note)
    strip_table_numbering(table)
    add_number_column(table)


_MINOR_WORDS = {"a", "an", "and", "as", "at", "but", "by", "for", "if", "in",
                "into", "nor", "of", "on", "onto", "or", "the", "to", "via",
                "vs", "with"}


def _title_segment(seg: str) -> str:
    words = seg.split(" ")
    n = len(words)
    out = []
    for i, w in enumerate(words):
        ai = next((k for k, ch in enumerate(w) if ch.isalpha()), None)
        if ai is None:                       # no letters (numbers, "&", etc.)
            out.append(w)
            continue
        is_minor = w[ai:].strip(".,()").lower() in _MINOR_WORDS
        if is_minor and 0 < i < n - 1:       # lowercase minor words mid-segment
            out.append(w[:ai] + w[ai:].lower())
        else:                                # capitalise first letter, keep the rest
            out.append(w[:ai] + w[ai].upper() + w[ai + 1:])
    return " ".join(out)


def _cap_words(s: str) -> str:
    """Title-case each ' - ' segment: first letter of each word capitalised,
    acronyms preserved, minor words (of, the, and, ...) kept lowercase unless
    first or last word of their segment."""
    return " - ".join(_title_segment(seg) for seg in s.split(" - "))


def safe_filename(s: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', "", s).strip()


# ----------------------------------------------------------------------------- #
# Main
# ----------------------------------------------------------------------------- #

H = {  # normalised header -> our key
    "name": "name", "customer": "customer", "course location": "location",
    "training contact": "contact", "training contact phone": "phone",
    "course date": "start", "course end date": "end", "other date": "other",
    "time (from)": "tfrom", "time (to)": "tto", "instructors": "instructor",
    "participant names": "participants",
    "recertification participants": "recert_participants",
}


def read_export(xlsx_path):
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    headers = {}
    for c in range(1, ws.max_column + 1):
        h = (ws.cell(row=1, column=c).value or "").strip().lower()
        if h in H:
            headers[H[h]] = c
    rows = []
    for r in range(2, ws.max_row + 1):
        if not any(ws.cell(row=r, column=c).value for c in range(1, ws.max_column + 1)):
            continue
        rows.append({k: ws.cell(row=r, column=col).value for k, col in headers.items()})
    return rows


def build_one(rec, template_map, out_dir, sort_alpha=True, include_cancelled=False):
    name_col = rec.get("name") or ""
    company = clean_company(rec.get("customer"))

    if not include_cancelled and "cancel" in name_col.lower():
        return ("skipped", "course marked CANCELLED", None, 0)

    ctype = detect_course_type(name_col)
    tpl = template_map.get(ctype)
    if not tpl:
        return ("skipped", f"no template for '{ctype}'", ctype, 0)

    contact_name = extract_contact_name(rec.get("contact"))
    info = {
        "Course Name": extract_course_name(name_col),
        "Company Name": company,
        "Location": extract_location(rec.get("location"), company, contact_name),
        "Contact": format_contact(contact_name, rec.get("phone")),
        "Date": format_dates(rec.get("start"), rec.get("other"), rec.get("end")),
        "Time": format_time(rec.get("tfrom"), rec.get("tto")),
        "Instructor": (rec.get("instructor") or "").strip(),
    }
    participants = parse_participants(rec.get("participants"),
                                     rec.get("recert_participants"), sort_alpha)

    doc = Document(tpl)
    fill_info_table(doc, info)
    fill_participants(doc, participants)

    start_label = _fmt_single(_first_date(rec.get("start"), rec.get("other"), rec.get("end")))
    instructor = info["Instructor"] or "Instructor"
    raw = " - ".join(p for p in (instructor, start_label, company) if p)
    fname = safe_filename(_cap_words(raw).rstrip(" .")) + ".docx"
    out_path = os.path.join(out_dir, fname)
    doc.save(out_path)
    return ("ok", out_path, ctype, len(participants))


def main():
    ap = argparse.ArgumentParser(description="Build course rosters from an Excel export.")
    ap.add_argument("export", help="Path to the Excel export (.xlsx)")
    ap.add_argument("--templates", default=".", help="Folder containing the blank templates")
    ap.add_argument("--out", default="rosters", help="Output folder")
    ap.add_argument("--keep-order", action="store_true",
                    help="Preserve export order instead of sorting names A-Z")
    ap.add_argument("--include-cancelled", action="store_true",
                    help="Also generate rosters for courses marked CANCELLED")
    args = ap.parse_args()

    os.makedirs(args.out, exist_ok=True)
    template_map = load_template_map(args.templates)
    if not template_map:
        sys.exit(f"No templates found in {args.templates}")
    print("Templates detected:")
    for t, p in template_map.items():
        print(f"  {t:9s} <- {os.path.basename(p)}")

    rows = read_export(args.export)
    print(f"\n{len(rows)} course(s) in export\n")
    made, skipped = 0, []
    for rec in rows:
        status, info, ctype, n = build_one(
            rec, template_map, args.out,
            sort_alpha=not args.keep_order,
            include_cancelled=args.include_cancelled)
        if status == "ok":
            made += 1
            print(f"  [{ctype:8s}] {n:2d} participants -> {os.path.basename(info)}")
        else:
            skipped.append((rec.get("customer"), info))
            print(f"  [SKIPPED ] {info}")
    print(f"\nDone: {made} roster(s) created"
          + (f", {len(skipped)} skipped" if skipped else ""))


if __name__ == "__main__":
    main()
